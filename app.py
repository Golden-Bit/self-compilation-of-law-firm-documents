import os
import time

import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO
import zipfile
from datetime import datetime
import locale
import random

from docx_to_pdf import convert_docx_to_pdf

# Configurazione della pagina
st.set_page_config(layout="wide")

# Imposta la localizzazione italiana
try:
    locale.setlocale(locale.LC_TIME, 'it_IT.UTF-8')
except locale.Error:
    st.warning("La localizzazione 'it_IT.UTF-8' non è disponibile. Verrà utilizzata la localizzazione predefinita.")

# Credenziali salvate in locale
USERNAME = "admin"
PASSWORD = "admin"


def standardize(txt):
    txt = str(txt).strip().removesuffix(".0")
    return txt


# Funzione di login
def login(username, password):
    return username == USERNAME and password == PASSWORD


# CSS personalizzato per nascondere le icone dei collegamenti
hide_link_icons_css = """
    <style>
    a {text-decoration: none;}
    .st-emotion-cache-gi0tri.e1nzilvr1 {display: none;}  /* Nasconde l'icona del collegamento nel titolo */
    </style>
"""
st.markdown(hide_link_icons_css, unsafe_allow_html=True)


def process_file(uploaded_file):
    file_type = uploaded_file.name.split('.')[-1]
    if file_type == 'csv':
        df = pd.read_csv(uploaded_file, delimiter=';').fillna(' ')
    elif file_type == 'xlsx':
        df = pd.read_excel(uploaded_file).fillna(' ')
    else:
        raise ValueError("Formato file non supportato. Caricare un file .csv o .xlsx.")
    return df


def extract_relevant_fields(df, soggetto, date_format):
    subset = df[df['Codice_Soggetto'] == soggetto].fillna(' ')
    contracts = "/".join(subset['CONTRATTO'].astype(str).unique())

    if date_format == "alfanumerico":
        data_odierna = datetime.now().strftime("%d %B %Y")
    elif date_format == "numerico":
        data_odierna = datetime.now().strftime("%d/%m/%Y")
    else:
        data_odierna = datetime.now().strftime("%d %B %Y")

    data_odierna = f" {data_odierna}"

    replacements = {
        '[Ragione_Sociale]': subset['Ragione_Sociale'].iloc[0],
        '[Indirizzo_Fornitura]': subset['Indirizzo_Fornitura'].iloc[0],
        '[Cap_Fornitura]': subset['Cap_Fornitura'].iloc[0],
        '[Comune_Fornitura]': subset['Comune_Fornitura'].iloc[0],
        '[Provincia_Fornitura]': subset['Provincia_Fornitura'].iloc[0],
        'IND_S': subset['Indirizzo_Spedizione'].iloc[0] if 'Indirizzo_Spedizione' in subset.columns else
        subset['Indirizzo_Fornitura'].iloc[0],
        'CAP_S': subset['Cap_Spedizione'].iloc[0] if 'Cap_Spedizione' in subset.columns else
        subset['Cap_Fornitura'].iloc[0],
        'COM_S': subset['Comune_Spedizione'].iloc[0] if 'Comune_Spedizione' in subset.columns else
        subset['Comune_Fornitura'].iloc[0],
        'PRO_S': subset['Provincia_Spedizione'].iloc[0] if 'Provincia_Spedizione' in subset.columns else
        subset['Provincia_Fornitura'].iloc[0],
        '[Email_Soggetto]': subset['Email_Soggetto'].iloc[0] if 'Email_Soggetto' in subset.columns else ' ',
        '[Pec_Soggetto]': subset['Pec_Soggetto'].iloc[0] if 'Pec_Soggetto' in subset.columns else ' ',
        '[CONTRATTO]': contracts,
        '[Codice_Commerciale]': subset['Codice_Commerciale'].iloc[0],
        '[Codice_Soggetto]': soggetto,
        '[Residuo ad oggi]': ' ',  # Placeholder
        'DATA': data_odierna,  # Data odierna formattata
    }

    return replacements


def get_identifier(row):

    codice_fiscale = row['Codice_Fiscale']
    partita_iva = row['Partita_Iva']
    codice_fiscale = standardize(codice_fiscale)
    #partita_iva = standardize(partita_iva)

    return row['Codice_Fiscale'] if codice_fiscale else row['Partita_Iva']


def extract_table_rows(identifier, df_fatture):
    fatture_contratto = df_fatture[
        (df_fatture['Codice fiscale'] == identifier) | (df_fatture['Partita IVA'] == identifier)].fillna(' ')

    table_rows = []
    for _, row in fatture_contratto.iterrows():
        table_row = {
            'Contratto': row['Contratto'],
            'Riferimento': row['N. documento'],
            'Data reg.': row['Data reg.'],
            'Scad.netto': row['Scad.netto'],
            'Importo sollecitabile': row['Importo sollecitabile']
        }
        table_rows.append(table_row)

    return table_rows


def get_residuo_ad_oggi(identifier, df_pratiche):
    pratica = df_pratiche[(df_pratiche['Codice fiscale'] == identifier) | (df_pratiche['Partita IVA'] == identifier)]
    if pratica.empty:
        raise ValueError("Identificativo non trovato nelle pratiche")
    residuo_ad_oggi = pratica['Residuo ad oggi'].fillna(' ').iloc[0]
    return residuo_ad_oggi


def get_numero_pratica(identifier, df_pratiche):
    pratica = df_pratiche[(df_pratiche['Codice fiscale'] == identifier) | (df_pratiche['Partita IVA'] == identifier)]
    if pratica.empty:
        raise ValueError("Identificativo non trovato nelle pratiche")
    numero_pratica = pratica['Numero pratica'].fillna(' ').iloc[0]
    return numero_pratica


def get_numero_affido(identifier, df_pratiche):
    pratica = df_pratiche[(df_pratiche['Codice fiscale'] == identifier) | (df_pratiche['Partita IVA'] == identifier)]
    if pratica.empty:
        raise ValueError("Identificativo non trovato nelle pratiche")
    numero_affido = pratica['Numero affido'].fillna(' ').iloc[0]
    return numero_affido


def replace_text(doc, replacements):
    for paragraph in doc.paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in paragraph.text:
                for run in paragraph.runs:
                    if old_text in run.text:
                        if old_text == "[Residuo ad oggi]":
                            run.text = run.text.replace(old_text, f"{new_text}")
                            run.bold = True  # Applicare il grassetto
                        else:
                            run.text = run.text.replace(old_text, str(new_text))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for old_text, new_text in replacements.items():
                        if old_text in paragraph.text:
                            for run in paragraph.runs:
                                if old_text in run.text:
                                    if old_text == "[Residuo ad oggi]":
                                        run.text = run.text.replace(old_text, f"{new_text}")
                                        run.bold = True  # Applicare il grassetto
                                    else:
                                        run.text = run.text.replace(old_text, str(new_text))


def update_document(input_path, replacements, table_rows):
    doc = Document(input_path)
    replace_text(doc, replacements)
    table = doc.tables[0]

    for row in table.rows[1:]:
        table._element.remove(row._element)

    for row_data in table_rows:
        row_cells = table.add_row().cells
        row_cells[0].text = row_data['Contratto']
        row_cells[1].text = row_data['Riferimento']
        row_cells[2].text = row_data['Data reg.']
        row_cells[3].text = row_data['Scad.netto']
        row_cells[4].text = row_data['Importo sollecitabile']

    output = BytesIO()
    doc.save(output)
    return output


# Analisi del documento per confermare i segnaposto
def analyze_document(doc_path):
    doc = Document(doc_path)
    placeholders = set()
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            placeholders.update([word for word in run.text.split() if word.startswith('[') and word.endswith(']')])
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        placeholders.update(
                            [word for word in run.text.split() if word.startswith('[') and word.endswith(']')])
    return placeholders


# Pagina di login separata
def show_login_page():
    st.title("Login")
    username = st.text_input("Username")
    password = st.text_input("Password", type="password")
    if st.button("Login"):
        if login(username, password):
            st.session_state['logged_in'] = True
            st.experimental_rerun()
        else:
            st.error("Username o password errati")


# Pagina principale
def show_main_page(doc_path):
    st.title('Studio Carotenuto - Portale Automazione')
    st.markdown("---")

    col1, col2, col3 = st.columns(3)
    with col1:
        uploaded_anagrafiche = st.file_uploader("Carica il file EXCEL 1 ANAGRAFICHE", type=["csv", "xlsx"])
    with col2:
        uploaded_fatture = st.file_uploader("Carica il file EXCEL 2 FATTURE", type=["csv", "xlsx"])
    with col3:
        uploaded_pratiche = st.file_uploader("Carica il file EXCEL 3 PRATICHE", type=["csv", "xlsx"])

    st.markdown("---")

    if uploaded_anagrafiche and uploaded_fatture and uploaded_pratiche:
        df_anagrafiche = process_file(uploaded_anagrafiche)
        df_fatture = process_file(uploaded_fatture)
        df_pratiche = process_file(uploaded_pratiche)

        date_format_option = st.selectbox('Seleziona il formato della data:', ['alfanumerico', 'numerico'])
        option = st.selectbox('Seleziona l\'opzione di generazione:', ['Singolo soggetto', 'Tutti i soggetti'])

        if option == 'Singolo soggetto':
            soggetti = df_anagrafiche['Codice_Soggetto'].unique()
            selected_soggetto = st.selectbox('Seleziona il Codice Soggetto', soggetti)
            st.markdown("---")

            if st.button('Genera Documento'):
                generate_single_document(selected_soggetto, df_anagrafiche, df_fatture, df_pratiche, doc_path,
                                         date_format_option)

        elif option == 'Tutti i soggetti':
            st.markdown("---")
            if st.button('Genera Documenti'):
                generate_all_documents(df_anagrafiche, df_fatture, df_pratiche, doc_path, date_format_option)


# Genera un singolo documento
def generate_single_document(soggetto, df_anagrafiche, df_fatture, df_pratiche, doc_path, date_format_option):
    try:
        replacements = extract_relevant_fields(df_anagrafiche, soggetto, date_format_option)
        identifier = get_identifier(df_anagrafiche[df_anagrafiche['Codice_Soggetto'] == soggetto].iloc[0])
        residuo_ad_oggi = get_residuo_ad_oggi(identifier, df_pratiche)
        numero_pratica = get_numero_pratica(identifier, df_pratiche)
        numero_affido = get_numero_affido(identifier, df_pratiche)
        replacements['[Residuo ad oggi]'] = residuo_ad_oggi
        table_rows = extract_table_rows(identifier, df_fatture)

        output = update_document(doc_path, replacements, table_rows)
        docx_filename = f"{replacements['[Ragione_Sociale]']}_{numero_affido}.docx"
        pdf_filename = f"{replacements['[Ragione_Sociale]']}_{numero_affido}.pdf"

        with open(docx_filename, "wb") as f:
            f.write(output.getvalue())

        convert_docx_to_pdf(docx_filename, pdf_filename)

        with open(pdf_filename, "rb") as f:
            pdf_output = f.read()

        st.download_button(
            label="Scarica il documento DOCX",
            data=output,
            file_name=docx_filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        st.download_button(
            label="Scarica il documento PDF",
            data=pdf_output,
            file_name=pdf_filename,
            mime="application/pdf"
        )

        os.remove(docx_filename)
        os.remove(pdf_filename)

    except Exception as e:
        st.error(f"Errore: {e}")


# Genera documenti per tutti i soggetti
def generate_all_documents(df_anagrafiche, df_fatture, df_pratiche, doc_path, date_format_option):
    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, 'a', zipfile.ZIP_DEFLATED) as zip_file:
        soggetti = df_anagrafiche['Codice_Soggetto'].unique()
        num_soggetti = len(soggetti)
        progress_bar = st.progress(0)

        for i, soggetto in enumerate(soggetti):
            try:

                time_1 = time.time()
                replacements = extract_relevant_fields(df_anagrafiche, soggetto, date_format_option)
                identifier = get_identifier(df_anagrafiche[df_anagrafiche['Codice_Soggetto'] == soggetto].iloc[0])
                print(identifier)
                residuo_ad_oggi = get_residuo_ad_oggi(identifier, df_pratiche)
                numero_pratica = get_numero_pratica(identifier, df_pratiche)
                numero_affido = get_numero_affido(identifier, df_pratiche)
                replacements['[Residuo ad oggi]'] = residuo_ad_oggi
                table_rows = extract_table_rows(identifier, df_fatture)
                time_2 = time.time()
                delta_1 = time_2 - time_1
                print(f"delta_1: {delta_1}")

                time_3 = time.time()
                output = update_document(doc_path, replacements, table_rows)
                time_4 = time.time()
                delta_2 = time_4 - time_3
                print(f"delta_2: {delta_2}")

                doc_filename = f"{replacements['[Ragione_Sociale]']}_{numero_affido}.docx"
                pdf_filename = f"{replacements['[Ragione_Sociale]']}_{numero_affido}.pdf"

                time_5 = time.time()
                with open(doc_filename, "wb") as f:
                    f.write(output.getvalue())
                time_6 = time.time()
                delta_3 = time_6 - time_5
                print(f"delta_3: {delta_3}")

                time_7 = time.time()
                convert_docx_to_pdf(doc_filename, pdf_filename)
                time_8 = time.time()
                delta_4 = time_8 - time_7
                print(f"delta_4: {delta_4}")

                time_9 = time.time()
                with open(doc_filename, "rb") as f:
                    zip_file.writestr(doc_filename, f.read())
                with open(pdf_filename, "rb") as f:
                    zip_file.writestr(pdf_filename, f.read())
                time_10 = time.time()
                delta_5 = time_10 - time_9
                print(f"delta_5: {delta_5}")

                os.remove(doc_filename)
                os.remove(pdf_filename)

            except Exception as e:
                st.error(f"Errore per il soggetto {soggetto}: {e}")

            progress_bar.progress((i + 1) / num_soggetti)

    zip_buffer.seek(0)
    st.download_button(
        label="Scarica tutti i documenti",
        data=zip_buffer,
        file_name="Lettere_Diffida.zip",
        mime="application/zip"
    )


# Percorso del documento
doc_path = "input_data/TemplateLetteraDiffida.docx"

# Gestione dello stato di login
if 'logged_in' not in st.session_state:
    show_login_page()
else:
    show_main_page(doc_path)
